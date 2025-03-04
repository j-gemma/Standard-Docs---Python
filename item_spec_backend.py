from openpyxl import load_workbook
from openpyxl import cell
from docx import Document
import re
from datetime import date
from tkinter import messagebox
from docxcompose.composer import Composer

# Class to represent an individual item with its details and formatted spec information
class AqItem:
    def __init__(self, item_no, mfr, model, category, spec, accs):
        
        self.item_no = item_no    # Item number
        self.mfr = mfr            # Manufacturer name
        self.model = model        # Model number
        self.category = category  # Category type of the item
        self.spec = spec          # Specification details of the item
        self.accs = accs          # List of accessories associated with the item
        
        # First line of the specification text
        self.line1 = "ITEM {} - {}\n".format(self.item_no.upper(), self.category.upper())
        
        # Second line for details (empty by default)
        self.line2 = ''         

# Class to represent an accessory with its details
class AqAcc:
    def __init__(self, mfr, model_no, spec):
        self.mfr = mfr            # Manufacturer of the accessory
        self.model_no = model_no  # Model number of the accessory
        self.spec = spec          # Specification details of the accessory


"""
    Main function to read an Excel file, process data, and save a Word document with item specifications.
    Args:
        xl_file (str): Path to the input Excel file.
        save_as (str): Desired name for the saved Word document.
        save_path (str): Directory to save the final document.
        to_compose (str): Whether to include front and back templates.
    Returns:
        list: List of item numbers with errors during processing.
"""
def main(xl_file, save_as, save_path, to_compose):

    TEMP_DIR_PATH = R"Q:\Standard Documents\Templates - Do not edit or move these files!\RUNTIME - TEMPORARY"
        
    #print('working')
    ws, column_dict = get_columns(xl_file)     # Load data and identify column indices
    items = make_spec_dict(ws, column_dict)    # Build a dictionary of items from the spreadsheet
    items, error_item_nos = clean_data(items)  # Clean and normalize data
    
    # Generate a specification document from the processed items
    item_spec = make_spec(items)
    
    #saved as temp because the doc may be modified again before saved with user input filename
    item_spec.save(R"Q:\Standard Documents\Templates - Do not edit or move these files!\RUNTIME - TEMPORARY" + R"\temp.docx")

    if to_compose != 'nocompose':
        item_spec = compose()  # Combine with front and back section templates if specified

    item_spec.save(save_path + "/" + save_as + '.docx')  
    return error_item_nos

#THIS DOES NOT WORK YET
"""
    Copy headers and footers from a source document to a target document.
    Args:
        header_doc (str): Path to the source document with the desired headers/footers.
        item_spec (Document): Target Word document to modify.
    Returns:
        Document: Updated target document with copied headers/footers.
"""
def copy_header_footer(header_doc, item_spec):
    
    source = Document(header_doc)
    target_header = item_spec.sections[0].header

    for paragraph in source.sections[0].header.paragraphs:
        new_para = target_header.add_paragraph(paragraph.text, paragraph.style)
        copy_format = paragraph.paragraph_format
        new_para_format = new_para.paragraph_format
        
        # Copy paragraph formatting
        new_para_format.alignment = copy_format.alignment
        new_para_format.first_line_indent = copy_format.first_line_indent
        new_para_format.keep_together = copy_format.keep_together
        new_para_format.keep_with_next = copy_format.keep_with_next
        new_para_format.left_indent = copy_format.left_indent
        new_para_format.line_spacing = copy_format.line_spacing
        new_para_format.line_spacing_rule = copy_format.line_spacing_rule
        new_para_format.page_break_before = copy_format.page_break_before
        new_para_format.right_indent = copy_format.right_indent
        new_para_format.space_after = copy_format.space_after
        new_para_format.space_before = copy_format.space_before
        new_para_format.widow_control = copy_format.widow_control

        # Add tab stops
        for tab_stop in copy_format.tab_stops:
            new_para_format.tab_stops.add_tab_stop(tab_stop.position, tab_stop.alignment)
    return item_spec


"""
    Create a Word document using a predefined template and fill it with item specifications.
    Args:
        items (dict): Dictionary of items and their details.
    Returns:
        Document: Word document populated with item specifications.
"""
def make_spec(items):
   
    ITEMIZED_SPEC = r"Q:\Standard Documents\Templates - Do not edit or move these files!\Specs\SECTION 114000 FOODSERVICE ITEM SPEC.docx"
    template = Document(ITEMIZED_SPEC)

    for key, item in items.items():
        # Add main item details
        if item.category == "OPEN NUMBER":
            para = template.add_paragraph(item.line1)
            para.style = template.styles['AqItem']
        else:
            para = template.add_paragraph(item.line1 + item.line2)
            para.style = template.styles['AqItem']
            
            para = template.add_paragraph(item.spec)
            para.style = template.styles['ItemSpec']

        # Add accessory details if present
        for acc in item.accs:
            para = template.add_paragraph(acc.spec)
            para.style = template.styles['Accessory']

    # Replace all occurrences of " Gc" with " GC" in text runs
    for paragraph in template.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(" Gc", " GC")

    return template


"""
    Clean and normalize item data to ensure consistency and process special cases.
    Args:
        items (dict): Dictionary of items to clean.
    Returns:
        tuple: (cleaned items, list of error item numbers)
"""
def clean_data(items):
    
    nikec_string = 'This item is not in the kitchen equipment contract.'
    #OLD_NIKEC_REGEX = r"NIKEC[\s/,]\s?((?:(?:BY|by|By)|(?:IN|in|In)|(?:EXISTING|existing|Existing))\s((?:[^/\s]+(?:/[^/\s]+)*)+))"
    NIKEC_REGEX = r"NIKEC\s*[/,\s]\s*(?:BY|IN|EXISTING)\s([^\n]*)"
    error_item_nos = []  # Track items with errors during processing

    for key, item in items.items():
        # Normalize manufacturer name
        if item.mfr == 'maxi':
            item.mfr = 'Maxi Movers'

        # Set line2 for standard items
        if item.category != 'OPEN NUMBER':
            item.line2 = '{} Model {}\n'.format(item.mfr, item.model)

        # Handle items with "NIKEC" in specifications
        if item.spec and item.spec.startswith("NIKEC"):
            punct = item.spec[5] if len(item.spec) >= 6 else ' '
            if punct == ',':
                punct = '/'

            match_obj = re.search(NIKEC_REGEX, item.spec, re.IGNORECASE)
            nikec_by = match_obj.group(1).title() if match_obj else 'FIX ME IN AUTOQUOTES'
            if not match_obj:
                error_item_nos.append(item.item_no)

            print(item.spec + '\n')
            print(match_obj.group(0) + '\n')
            item.line2 = match_obj.group(0).upper() + '\n'
            print(item.line2)
            item.spec = nikec_string
            item.accs = []

        # Normalize accessory details
        for acc in item.accs:
            if acc.mfr == item.mfr:
                acc.spec = "Model {} {}".format(acc.model_no, acc.spec) if acc.model_no else acc.spec
            elif acc.mfr:
                acc.spec = "{} Model {} {}".format(acc.mfr, acc.model_no, acc.spec) if acc.model_no else acc.spec

    return items, error_item_nos

"""
    Combine front, back, and middle document templates into a single document.
    Returns:
        Document: Combined Word document.
"""
def compose():
    
    front = r"Q:\Standard Documents\Templates - Do not edit or move these files!\Specs\GENERAL SPEC SOURCE FRONT.docx"
    back = r"Q:\Standard Documents\Templates - Do not edit or move these files!\Specs\GENERAL SPEC SOURCE BACK.docx"
    middle = r"Q:\Standard Documents\Templates - Do not edit or move these files!\RUNTIME - TEMPORARY\temp.docx"

    composer = Composer(Document(front))
    composer.append(Document(middle))
    composer.append(Document(back))

    return composer.doc


"""
    Identify column indices for required fields in the Excel spreadsheet.
    Args:
        filename (str): Path to the Excel file.
    Returns:
        tuple: Worksheet and dictionary mapping field names to column letters.
"""
def get_columns(filename):
    
    columns = {}
    find_columns = ['ItemNo', 'Mfr', 'Model', 'Category', 'Spec']
    wb = load_workbook(filename=filename)
    ws = wb['Spreadsheet']

    for column_name in find_columns:
        for cell in ws[1]:  # Header row
            if cell.value == column_name:
                columns[column_name] = cell.column_letter

    return ws, columns

"""
    Build a dictionary of items and their details from the spreadsheet data.
    Args:
        ws (Worksheet): Worksheet object.
        column_dict (dict): Column mappings from field names to letters.
    Returns:
        dict: Dictionary of AqItem objects indexed by item numbers.
"""
def make_spec_dict(ws, column_dict):
    
    spec_dict = {}
    item_no_col = ws[column_dict['ItemNo']]

    for row_num in range(1, len(item_no_col)):
        row = str(row_num + 1)
        item_cell = item_no_col[row_num]

        if item_cell.value and item_cell.fill.fgColor.rgb in ['FFFFFFFF', '00FFFFFF']:
            # Extract primary item data
            item_no = item_cell.value
            mfr = ws[column_dict['Mfr'] + row].value
            model = ws[column_dict['Model'] + row].value
            category = ws[column_dict['Category'] + row].value
            spec = ws[column_dict['Spec'] + row].value
            accs = []

            # Collect associated accessory rows
            tempidx = row_num + 1
            while tempidx < len(item_no_col) and item_no_col[tempidx].value is None:
                acc_row = str(tempidx + 1)
                if item_no_col[tempidx].fill.fgColor.rgb in ['FFFAFAD2', '00FAFAD2']:
                    accs.append(AqAcc(
                        ws[column_dict['Mfr'] + acc_row].value,
                        ws[column_dict['Model'] + acc_row].value,
                        ws[column_dict['Spec'] + acc_row].value
                    ))
                tempidx += 1

            # Add primary item to the specification dictionary
            spec_dict[item_no] = AqItem(item_no, mfr, model, category, spec, accs)

    return spec_dict
